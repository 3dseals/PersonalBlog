from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_style(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 153)
    return heading

def add_paragraph_text(doc, text, bold=False, size=11):
    """添加段落文本"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    return p

def create_info_table(doc, data):
    """创建信息表格"""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    for i, (key, value) in enumerate(data):
        row = table.rows[i]
        cell1 = row.cells[0]
        cell2 = row.cells[1]
        cell1.text = key
        cell2.text = value
        set_cell_shading(cell1, 'E8F4F8')
        for cell in [cell1, cell2]:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)
    return table

def add_skill_bar(doc, skill, level):
    """添加技能条"""
    p = doc.add_paragraph()
    run = p.add_run(f"{skill}：{'█' * level}{'░' * (10-level)}")
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

def main():
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # 标题
    title = doc.add_heading('王子亮 | 研发技术总监', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('15年+全栈研发经验 · 游戏/AI社交领域 · 团队管理与技术架构')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # 空行
    
    # 基本信息
    add_heading_style(doc, '📋 基本信息', 1)
    basic_info = [
        ('姓名', '王子亮 / 男 / 1987'),
        ('学历', '本科 211 · 哈尔滨工程大学 · 计算机科学与技术'),
        ('工作年限', '15年+'),
        ('期望职位', '研发类技术总监'),
        ('期望薪资', '30K'),
        ('期望城市', '成都'),
    ]
    create_info_table(doc, basic_info)
    
    doc.add_paragraph()
    
    # 联系方式
    add_heading_style(doc, '📞 联系方式', 1)
    contact_info = [
        ('手机', '15802857662'),
        ('Email', 'lional.king@gmail.com'),
        ('QQ', '870966629'),
        ('技术博客', 'https://ibunnyteam.github.io'),
        ('Github', 'https://github.com/3dseals'),
    ]
    create_info_table(doc, contact_info)
    
    doc.add_paragraph()
    
    # 核心优势
    add_heading_style(doc, '💡 核心优势', 1)
    
    add_paragraph_text(doc, '🎯 全栈技术覆盖', bold=True, size=12)
    advantages1 = [
        '• 客户端：Unity3D / Cocos2d-x / Cocos Creator / React.js',
        '• 服务端：Node.js / Go / Java / Skynet / C++',
        '• 数据层：MySQL / Redis / MongoDB / 高并发IO设计',
        '• DevOps：Jenkins / Docker / Git / 自动化部署',
    ]
    for item in advantages1:
        add_paragraph_text(doc, item, size=10)
    
    add_paragraph_text(doc, '🤝 团队协同能力', bold=True, size=12)
    advantages2 = [
        '• 主导多个从0到1的项目架构设计与团队搭建',
        '• 建立规范化开发流程（OKR/Redmine/代码审查）',
        '• 跨部门协作经验：策划、美术、运营、测试',
        '• AI辅助编程实践，提升团队研发效率',
    ]
    for item in advantages2:
        add_paragraph_text(doc, item, size=10)
    
    add_paragraph_text(doc, '📊 业务成果', bold=True, size=12)
    advantages3 = [
        '• 累计参与项目流水超 3000万+',
        '• 主导多款游戏从研发到上线全流程',
        '• 具备游戏运营思维，理解商业化设计',
    ]
    for item in advantages3:
        add_paragraph_text(doc, item, size=10)
    
    doc.add_paragraph()
    
    # 技术栈
    add_heading_style(doc, '🛠️ 技术栈', 1)
    
    tech_categories = [
        ('客户端开发', [
            ('Unity3D', 10), ('Cocos2d-x', 10), ('Cocos Creator', 8), ('React.js', 6)
        ]),
        ('服务端开发', [
            ('Node.js', 10), ('Go', 8), ('Java', 8), ('Spring Cloud', 8), ('Skynet/Lua', 8), ('C++', 6)
        ]),
        ('数据库 & 中间件', [
            ('MySQL', 10), ('Redis', 10), ('MongoDB', 8), ('消息队列(MQ)', 8)
        ]),
        ('工具链 & DevOps', [
            ('Git/SVN', 10), ('Jenkins', 8), ('Docker', 6)
        ]),
    ]
    
    for category, skills in tech_categories:
        add_paragraph_text(doc, category, bold=True, size=11)
        for skill, level in skills:
            add_skill_bar(doc, skill, level)
    
    doc.add_paragraph()
    
    # 工作经历
    add_heading_style(doc, '💼 工作经历', 1)
    
    jobs = [
        {
            'period': '2024.07 - 至今',
            'company': '四川循东科技有限公司',
            'position': '技术负责人',
            'desc': '专注于AI驱动的虚拟社交和数字娱乐领域的创新科技公司',
            'duties': [
                '负责技术团队管理，建立规范化开发流程与管理机制',
                '推动AI辅助编程实践（Cursor等），提升团队研发效率',
                '主导产品技术架构设计与技术选型',
            ]
        },
        {
            'period': '2023.03 - 2024.06',
            'company': '麓卓互动科技有限公司（成都分公司）',
            'position': '服务端技术负责人',
            'desc': '专注于AI社交和元宇宙领域，打造3D虚拟社交应用《Party Yoo》',
            'duties': [
                '负责AI社交产品服务端架构设计与开发',
                '参与3D虚拟社交应用的游戏模块开发',
                '推行OKR绩效管理，优化团队协作流程',
                '产品已在印尼等国际市场内测上线',
            ]
        },
        {
            'period': '2013.05 - 2023.02',
            'company': '成都黑骑士网络科技（创业）',
            'position': '技术合伙人 / 技术总监',
            'desc': '创业型手游公司，拥有多款上线游戏',
            'duties': [
                '从0到1搭建技术团队，建立完整开发流程',
                '主导多款游戏全栈开发（客户端+服务端）',
                '负责技术架构设计、性能优化、线上运维',
                '参与游戏运营策略制定，理解商业化设计',
                '核心项目流水累计超2000万',
            ]
        },
        {
            'period': '2011.12 - 2013.04',
            'company': '上海慕和网络',
            'position': '游戏开发工程师',
            'desc': '手游公司，2013年被凤凰传媒收购，月流水达1500万',
            'duties': [
                '参与多款SLG游戏开发',
                '参与游戏引擎研发与平台SDK研发',
                '见证公司从50人发展到400人的高速成长期',
            ]
        },
        {
            'period': '2010.07 - 2011.11',
            'company': 'TCL通讯',
            'position': 'Android开发工程师',
            'desc': '专注于Android手机研发，与阿尔卡特、黑莓等厂商合作',
            'duties': [
                '参与Android系统上层应用定制开发',
                '负责Contacts、Launcher等核心模块开发',
                '深入理解Android底层机制，为后续移动开发奠定基础',
            ]
        },
    ]
    
    for job in jobs:
        p = doc.add_paragraph()
        run = p.add_run(f"{job['period']} | {job['company']}")
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153)
        
        add_paragraph_text(doc, f"职位：{job['position']}", size=10)
        add_paragraph_text(doc, f"公司简介：{job['desc']}", size=10)
        add_paragraph_text(doc, '工作内容：', bold=True, size=10)
        for duty in job['duties']:
            add_paragraph_text(doc, f"• {duty}", size=10)
        doc.add_paragraph()
    
    # 项目经历
    add_heading_style(doc, '🎮 项目经历', 1)
    
    projects = [
        {
            'year': '2020',
            'name': '无尽之界',
            'type': '挂机对战冒险游戏',
            'role': '技术负责人',
            'tech': [
                '客户端：Unity3D + 战斗模块 + 逐帧验证 + 热更新',
                '服务端：Node.js分布式架构 + Redis高并发IO + RPC通信',
                '特色：灯光Shader、网络协议优化、战斗服/挂机合成设计',
            ]
        },
        {
            'year': '2019',
            'name': '三国诸侯纷争',
            'type': 'SLG策略游戏',
            'role': '全栈开发',
            'revenue': '200万+/半年',
            'tech': [
                '客户端：Unity3D + ILRuntime热更新',
                '服务端：Java NIO + Protobuf协议',
                '职责：多平台对接、支付系统、后台功能开发',
            ]
        },
        {
            'year': '2017',
            'name': 'Survive 活下去',
            'type': '丧尸生存游戏',
            'role': '技术负责人',
            'revenue': '400万+/年',
            'tech': [
                '首次采用JS作为业务逻辑语言',
                '多服架构 + HTTP/Socket双通信',
                '全球唯一服 + 消息推送系统',
                '建立完整工具链与开发流程（Redmine工作流）',
            ]
        },
        {
            'year': '2016',
            'name': '沙巴克传奇',
            'type': '传奇类RPG',
            'role': '全栈开发',
            'revenue': '500万+/半年',
            'tech': [
                '客户端：Cocos2d-Lua + 热更新 + 逐帧动画',
                '服务端：C++ + Lua热更新（活动系统）',
                '后台：PHP支付登录平台，解耦业务逻辑',
            ]
        },
        {
            'year': '2014',
            'name': '魔卡世纪',
            'type': '即时对战卡牌',
            'role': '客户端主程',
            'tech': [
                'Cocos2d-x (C++) 早期版本',
                '解决大量引擎底层问题：动画、灰度图、输入框、热更新',
                '加拿大AppStore上线，首款原创游戏',
            ]
        },
    ]
    
    for proj in projects:
        p = doc.add_paragraph()
        title_text = f"{proj['year']} | {proj['name']}"
        run = p.add_run(title_text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153)
        
        info_text = f"类型：{proj['type']} | 角色：{proj['role']}"
        if 'revenue' in proj:
            info_text += f" | 流水：{proj['revenue']}"
        add_paragraph_text(doc, info_text, size=10)
        
        add_paragraph_text(doc, '技术架构：', bold=True, size=10)
        for tech in proj['tech']:
            add_paragraph_text(doc, f"• {tech}", size=10)
        doc.add_paragraph()
    
    # 技术总结
    add_heading_style(doc, '📚 技术总结', 1)
    
    tech_summary = [
        ('游戏引擎', 'Unity3D / Cocos2d-x / Cocos2d-Lua / Cocos Creator'),
        ('服务端框架', 'Skynet / Pinus / ThinkJS / Go / Spring Cloud'),
        ('微服务组件', 'Nacos / Gateway / Feign / XXL-Job / Seata'),
        ('消息队列', 'RabbitMQ / Kafka / RocketMQ'),
        ('Web开发', 'React.js / Node.js / PHP / HTML5'),
        ('数据库', 'MySQL / Redis / MongoDB / Memcached'),
        ('协议 & 通信', 'Protobuf / WebSocket / HTTP / RPC'),
        ('工具链', 'Webpack / Gulp / Python脚本'),
        ('DevOps', 'Git / SVN / Jenkins / Redmine / Doxygen'),
        ('平台对接', 'AppStore / Google Play / 微信 / 微博 / 各类SDK'),
    ]
    create_info_table(doc, tech_summary)
    
    doc.add_paragraph()
    
    # 致谢
    add_heading_style(doc, '🙏 致谢', 1)
    
    thanks_text = '''感谢您花时间阅读我的简历。

15年的技术积累让我具备了从客户端到服务端、从架构设计到团队管理的全栈能力。我热爱技术，享受解决复杂问题的过程，也乐于带领团队一起成长。

期待能有机会与您共事，共同创造价值。'''
    
    p = doc.add_paragraph()
    run = p.add_run(thanks_text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    # 保存文档
    doc.save('王子亮_研发技术总监_简历.docx')
    print('简历已生成：王子亮_研发技术总监_简历.docx')

if __name__ == '__main__':
    main()
